"""Generate professional DOCX and PDF from PROJECT_DOCUMENTATION.md."""

import re
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
BASE = Path(__file__).parent
MD_PATH = BASE / "PROJECT_DOCUMENTATION.md"
DOCX_PATH = BASE / "PROJECT_DOCUMENTATION.docx"
PDF_PATH = BASE / "PROJECT_DOCUMENTATION.pdf"

AUTHOR = "Saikethana"
PROJECT_NAME = "Banking Application"
DATE_STR = date.today().strftime("%B %d, %Y")

FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"
FONT_CODE = "Consolas"

COLOR_PRIMARY = RGBColor(26, 26, 46)       # #1a1a2e
COLOR_ACCENT = RGBColor(59, 130, 246)       # #3b82f6
COLOR_TABLE_HEADER = RGBColor(26, 26, 46)
COLOR_TABLE_HEADER_TEXT = RGBColor(255, 255, 255)
COLOR_TABLE_ALT = RGBColor(240, 242, 245)
COLOR_CODE_BG = RGBColor(245, 245, 250)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right with (sz, color)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, (sz, color) in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{sz}" '
            f'w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_page_number(section):
    """Add page number to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)
    for r in [run, run2, run3]:
        r.font.size = Pt(9)
        r.font.name = FONT_BODY
        r.font.color.rgb = RGBColor(128, 128, 128)


def style_table(table):
    """Apply professional styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Style header row
    if table.rows:
        header_row = table.rows[0]
        for cell in header_row.cells:
            set_cell_shading(cell, "1A1A2E")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.color.rgb = COLOR_TABLE_HEADER_TEXT
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.name = FONT_BODY
        # Style body rows
        for i, row in enumerate(table.rows[1:], 1):
            for cell in row.cells:
                if i % 2 == 0:
                    set_cell_shading(cell, "F0F2F5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = FONT_BODY


# ---------------------------------------------------------------------------
# Markdown Parser
# ---------------------------------------------------------------------------

def parse_md(text: str):
    """Parse markdown into a list of block elements."""
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules
        if re.match(r"^---+\s*$", line):
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            blocks.append(("heading", level, m.group(2).strip()))
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            lang = line.strip().lstrip("`").strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append(("code", "\n".join(code_lines)))
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[-\s|:]+\|\s*$", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                if not re.match(r"^\s*\|[-\s|:]+\|\s*$", lines[i]):
                    table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                rows.append(cells)
            blocks.append(("table", rows))
            continue

        # Bullet list item
        m = re.match(r"^(\s*)[*-]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            level = indent // 2
            blocks.append(("bullet", level, m.group(2).strip()))
            i += 1
            continue

        # Numbered list item
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            level = indent // 3
            blocks.append(("numbered", level, m.group(2).strip()))
            i += 1
            continue

        # Paragraph text
        if line.strip():
            blocks.append(("para", line.strip()))
            i += 1
            continue

        i += 1

    return blocks


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

def add_formatted_text(paragraph, text):
    """Add text with bold, italic, inline code, and link formatting."""
    # Pattern: **bold**, *italic*, `code`, [text](url)
    pattern = r"(\*\*(.+?)\*\*|`([^`]+)`|\*([^*]+)\*|\[([^\]]+)\]\([^)]+\))"
    pos = 0
    for m in re.finditer(pattern, text):
        # Add text before match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = FONT_BODY
            run.font.size = Pt(10.5)

        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.font.bold = True
            run.font.name = FONT_BODY
            run.font.size = Pt(10.5)
        elif m.group(3):  # inline code
            run = paragraph.add_run(m.group(3))
            run.font.name = FONT_CODE
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif m.group(4):  # italic
            run = paragraph.add_run(m.group(4))
            run.font.italic = True
            run.font.name = FONT_BODY
            run.font.size = Pt(10.5)
        elif m.group(5):  # link text
            run = paragraph.add_run(m.group(5))
            run.font.color.rgb = COLOR_ACCENT
            run.font.name = FONT_BODY
            run.font.size = Pt(10.5)

        pos = m.end()

    # Remaining text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = FONT_BODY
        run.font.size = Pt(10.5)


# ---------------------------------------------------------------------------
# DOCX Builder
# ---------------------------------------------------------------------------

def build_docx():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Page numbers
    add_page_number(section)

    # -- Default styles --
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ==============================
    # TITLE PAGE
    # ==============================
    for _ in range(6):
        doc.add_paragraph("")

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = COLOR_ACCENT
    run.font.size = Pt(14)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(PROJECT_NAME)
    run.font.name = FONT_HEADING
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Project Documentation")
    run.font.name = FONT_HEADING
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = COLOR_ACCENT
    run.font.size = Pt(14)

    # Author
    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Author")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.name = FONT_BODY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(AUTHOR)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = FONT_BODY
    run.font.color.rgb = COLOR_PRIMARY

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(DATE_STR)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = FONT_BODY

    # ==============================
    # TABLE OF CONTENTS PAGE
    # ==============================
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Table of Contents")
    run.font.name = FONT_HEADING
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    # TOC field code — Word will populate this when user presses "Update Fields"
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run("(Right-click → Update Field to generate Table of Contents)")
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(150, 150, 150)
    run4.font.italic = True

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    # ==============================
    # DOCUMENT CONTENT
    # ==============================
    doc.add_page_break()

    md_text = MD_PATH.read_text(encoding="utf-8")

    # Remove the markdown title (first # line) and TOC section — we handle these in title page
    md_text = re.sub(r"^# Banking Application.*?\n", "", md_text, count=1)
    md_text = re.sub(
        r"## Table of Contents\n.*?(?=\n## )", "", md_text, flags=re.DOTALL, count=1
    )

    blocks = parse_md(md_text)
    numbered_counter = 0

    for block in blocks:
        btype = block[0]

        if btype == "heading":
            level = block[1]
            text = block[2]

            # Remove markdown anchor links like [text](#anchor)
            text = re.sub(r"\[([^\]]+)\]\(#[^)]+\)", r"\1", text)

            if level == 2:
                doc.add_page_break()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                # Add a colored bar before section headings
                run = p.add_run("▎ ")
                run.font.color.rgb = COLOR_ACCENT
                run.font.size = Pt(22)
                run = p.add_run(text)
                run.font.name = FONT_HEADING
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = COLOR_PRIMARY
                # Register as Heading 1 for TOC
                p.style = doc.styles["Heading 1"]
                for r in p.runs:
                    r.font.name = FONT_HEADING
                    r.font.color.rgb = COLOR_PRIMARY

            elif level == 3:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(text)
                run.font.name = FONT_HEADING
                run.font.size = Pt(15)
                run.font.bold = True
                run.font.color.rgb = COLOR_PRIMARY
                p.style = doc.styles["Heading 2"]
                for r in p.runs:
                    r.font.name = FONT_HEADING
                    r.font.color.rgb = COLOR_PRIMARY

            elif level == 4:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(text)
                run.font.name = FONT_HEADING
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(60, 60, 80)
                p.style = doc.styles["Heading 3"]
                for r in p.runs:
                    r.font.name = FONT_HEADING

        elif btype == "para":
            text = block[1]
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, text)

        elif btype == "bullet":
            level = block[1]
            text = block[2]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
            p.paragraph_format.space_after = Pt(3)
            # Bullet character
            run = p.add_run("• " if level == 0 else "◦ ")
            run.font.name = FONT_BODY
            run.font.size = Pt(10.5)
            add_formatted_text(p, text)

        elif btype == "numbered":
            level = block[1]
            text = block[2]
            numbered_counter += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
            p.paragraph_format.space_after = Pt(3)
            # Extract leading number if present in text
            m = re.match(r"^(\d+)\.\s*(.*)", text)
            if m:
                text = m.group(2)
            add_formatted_text(p, text)

        elif btype == "code":
            code_text = block[1]
            for code_line in code_text.split("\n"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(code_line if code_line else " ")
                run.font.name = FONT_CODE
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(50, 50, 60)

            # Add small spacer after code block
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)

        elif btype == "table":
            rows = block[1]
            if not rows or len(rows) < 2:
                continue

            num_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set column widths based on content
            total_width = Inches(6.0)

            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < num_cols:
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(3)
                        p.paragraph_format.space_after = Pt(3)

                        # Handle inline code in table cells
                        parts = re.split(r"(`[^`]+`)", cell_text)
                        for part in parts:
                            if part.startswith("`") and part.endswith("`"):
                                run = p.add_run(part[1:-1])
                                run.font.name = FONT_CODE
                                run.font.size = Pt(8.5)
                                run.font.color.rgb = RGBColor(80, 80, 80)
                            else:
                                # Handle bold
                                bold_parts = re.split(r"(\*\*[^*]+\*\*)", part)
                                for bp in bold_parts:
                                    if bp.startswith("**") and bp.endswith("**"):
                                        run = p.add_run(bp[2:-2])
                                        run.font.bold = True
                                        run.font.name = FONT_BODY
                                        run.font.size = Pt(9)
                                    else:
                                        run = p.add_run(bp)
                                        run.font.name = FONT_BODY
                                        run.font.size = Pt(9)

            style_table(table)

            # Spacer after table
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

    # ==============================
    # Save
    # ==============================
    doc.save(str(DOCX_PATH))
    print(f"DOCX saved: {DOCX_PATH}")
    return doc


# ---------------------------------------------------------------------------
# PDF Builder (using fpdf2)
# ---------------------------------------------------------------------------

from fpdf import FPDF

class BankingPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        if self.page_no() > 1:
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"{PROJECT_NAME} - Project Documentation", align="L")
            self.ln(4)
            self.set_draw_color(220, 220, 220)
            self.line(10, self.get_y(), self.w - 10, self.get_y())
            self.ln(6)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "", 8)
        self.set_text_color(150, 150, 150)
        if self.page_no() > 1:
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    def add_title_page(self):
        self.add_page()
        self.ln(60)

        # Decorative line
        self.set_draw_color(59, 130, 246)
        self.set_line_width(0.8)
        self.line(50, self.get_y(), self.w - 50, self.get_y())
        self.ln(12)

        # Title
        self.set_font("Helvetica", "B", 32)
        self.set_text_color(26, 26, 46)
        self.cell(0, 16, PROJECT_NAME, align="C")
        self.ln(14)

        # Subtitle
        self.set_font("Helvetica", "", 16)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, "Project Documentation", align="C")
        self.ln(14)

        # Decorative line
        self.line(50, self.get_y(), self.w - 50, self.get_y())
        self.ln(30)

        # Author
        self.set_font("Helvetica", "", 11)
        self.set_text_color(120, 120, 120)
        self.cell(0, 8, "Author", align="C")
        self.ln(8)
        self.set_font("Helvetica", "B", 16)
        self.set_text_color(26, 26, 46)
        self.cell(0, 10, AUTHOR, align="C")
        self.ln(16)

        # Date
        self.set_font("Helvetica", "", 12)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, DATE_STR, align="C")

    def add_toc_page(self, toc_items):
        self.add_page()
        self.set_font("Helvetica", "B", 22)
        self.set_text_color(26, 26, 46)
        self.cell(0, 14, "Table of Contents", align="L")
        self.ln(14)

        for level, title, page in toc_items:
            indent = 4 + (level - 2) * 8
            if level == 2:
                self.set_font("Helvetica", "B", 11)
                self.set_text_color(26, 26, 46)
            elif level == 3:
                self.set_font("Helvetica", "", 10)
                self.set_text_color(60, 60, 80)
            else:
                self.set_font("Helvetica", "", 9)
                self.set_text_color(80, 80, 100)

            self.set_x(10 + indent)
            # Title
            title_width = self.w - 20 - indent - 15
            self.cell(title_width, 7, title, align="L")
            self.cell(15, 7, str(page), align="R")
            self.ln(7)


def clean_md_text(text):
    """Remove markdown formatting for PDF plain text."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Replace special characters that may not be in latin-1
    text = text.replace("→", "->")
    text = text.replace("₹", "Rs.")
    text = text.replace("━", "-")
    text = text.replace("▎", "|")
    text = text.replace("▶", ">")
    text = text.replace("◦", "o")
    text = text.replace("·", ".")
    text = text.replace("—", "--")
    text = text.replace("–", "-")
    text = text.replace("•", "-")
    text = text.replace("\u2018", "'")
    text = text.replace("\u2019", "'")
    text = text.replace("\u201c", '"')
    text = text.replace("\u201d", '"')
    text = text.replace("\u2264", "<=")  # ≤
    text = text.replace("\u2265", ">=")  # ≥
    text = text.replace("\u2014", "--")  # em dash
    text = text.replace("\u2013", "-")   # en dash
    text = text.replace("\u00b7", ".")   # middle dot
    text = text.replace("\u2026", "...")  # ellipsis
    text = text.replace("\u00a0", " ")   # non-breaking space
    # Strip any remaining non-latin-1 chars
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text


def build_pdf():
    md_text = MD_PATH.read_text(encoding="utf-8")

    # Remove title and TOC from md
    md_text = re.sub(r"^# Banking Application.*?\n", "", md_text, count=1)
    md_text = re.sub(
        r"## Table of Contents\n.*?(?=\n## )", "", md_text, flags=re.DOTALL, count=1
    )

    blocks = parse_md(md_text)

    # First pass: collect TOC items and estimate pages
    # We'll do a two-pass approach: generate once, collect pages, then regenerate with TOC

    def render_content(pdf, blocks, collect_toc=False):
        toc_items = []
        for block in blocks:
            btype = block[0]

            if btype == "heading":
                level = block[1]
                text = clean_md_text(block[2])
                text = re.sub(r"\[([^\]]+)\]\(#[^)]+\)", r"\1", text)

                if level == 2:
                    pdf.add_page()
                    if collect_toc:
                        toc_items.append((level, text, pdf.page_no()))
                    pdf.set_font("Helvetica", "B", 20)
                    pdf.set_text_color(26, 26, 46)
                    # Accent bar
                    pdf.set_fill_color(59, 130, 246)
                    pdf.rect(10, pdf.get_y(), 3, 10, "F")
                    pdf.set_x(17)
                    pdf.cell(0, 12, text, align="L")
                    pdf.ln(16)

                elif level == 3:
                    if collect_toc:
                        toc_items.append((level, text, pdf.page_no()))
                    pdf.ln(4)
                    pdf.set_font("Helvetica", "B", 14)
                    pdf.set_text_color(26, 26, 46)
                    pdf.cell(0, 10, text, align="L")
                    pdf.ln(10)

                elif level == 4:
                    if collect_toc:
                        toc_items.append((level, text, pdf.page_no()))
                    pdf.ln(2)
                    pdf.set_font("Helvetica", "B", 11)
                    pdf.set_text_color(60, 60, 80)
                    pdf.cell(0, 8, text, align="L")
                    pdf.ln(8)

            elif btype == "para":
                text = clean_md_text(block[1])
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 40)
                pdf.multi_cell(0, 5.5, text)
                pdf.ln(3)

            elif btype == "bullet":
                level = block[1]
                text = clean_md_text(block[2])
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 40)
                indent = 14 + level * 8
                pdf.set_x(indent)
                bullet = "-  " if level == 0 else "o  "
                pdf.multi_cell(pdf.w - indent - 10, 5.5, bullet + text)
                pdf.ln(1)

            elif btype == "numbered":
                level = block[1]
                text = clean_md_text(block[2])
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 40)
                indent = 14 + level * 8
                pdf.set_x(indent)
                pdf.multi_cell(pdf.w - indent - 10, 5.5, text)
                pdf.ln(1)

            elif btype == "code":
                code_text = clean_md_text(block[1])
                pdf.ln(2)
                pdf.set_fill_color(245, 245, 250)
                pdf.set_font("Courier", "", 8)
                pdf.set_text_color(50, 50, 60)
                for code_line in code_text.split("\n"):
                    if not code_line:
                        code_line = " "
                    pdf.set_x(14)
                    pdf.cell(pdf.w - 28, 4.5, code_line, fill=True)
                    pdf.ln(4.5)
                pdf.ln(4)

            elif btype == "table":
                rows = block[1]
                if not rows or len(rows) < 2:
                    continue

                num_cols = len(rows[0])
                available_width = pdf.w - 20
                col_width = available_width / num_cols

                # Calculate column widths based on content
                col_widths = []
                for c in range(num_cols):
                    max_len = 0
                    for row in rows:
                        if c < len(row):
                            max_len = max(max_len, len(clean_md_text(row[c])))
                    col_widths.append(max(max_len, 5))
                total = sum(col_widths)
                col_widths = [w / total * available_width for w in col_widths]
                # Cap max width
                col_widths = [min(w, available_width * 0.4) for w in col_widths]
                # Redistribute
                remaining = available_width - sum(col_widths)
                if remaining > 0:
                    for i in range(len(col_widths)):
                        col_widths[i] += remaining / len(col_widths)

                pdf.ln(2)
                for r_idx, row_data in enumerate(rows):
                    if r_idx == 0:
                        pdf.set_fill_color(26, 26, 46)
                        pdf.set_text_color(255, 255, 255)
                        pdf.set_font("Helvetica", "B", 8)
                    else:
                        if r_idx % 2 == 0:
                            pdf.set_fill_color(240, 242, 245)
                        else:
                            pdf.set_fill_color(255, 255, 255)
                        pdf.set_text_color(40, 40, 40)
                        pdf.set_font("Helvetica", "", 8)

                    x_start = pdf.get_x()
                    max_h = 5
                    # Calculate row height
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < num_cols:
                            text = clean_md_text(cell_text)
                            w = col_widths[c_idx]
                            lines = pdf.multi_cell(w, 5, text, split_only=True)
                            h = max(len(lines) * 5, 5)
                            max_h = max(max_h, h)

                    # Check if we need a page break
                    if pdf.get_y() + max_h > pdf.h - 25:
                        pdf.add_page()

                    y_start = pdf.get_y()
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < num_cols:
                            text = clean_md_text(cell_text)
                            w = col_widths[c_idx]
                            x = 10 + sum(col_widths[:c_idx])
                            pdf.set_xy(x, y_start)
                            # Draw cell background
                            pdf.rect(x, y_start, w, max_h, "F")
                            pdf.set_xy(x + 1, y_start + 1)
                            pdf.multi_cell(w - 2, 5, text)

                    pdf.set_y(y_start + max_h)

                pdf.ln(6)

        return toc_items

    # Pass 1: render to collect TOC page numbers
    pdf1 = BankingPDF()
    pdf1.add_title_page()
    pdf1.add_page()  # placeholder TOC page
    toc_items = render_content(pdf1, blocks, collect_toc=True)

    # Adjust page numbers: TOC takes 1 page (page 2), content starts at page 3
    # Page offset already handled since we rendered with title + TOC placeholder

    # Pass 2: final render with TOC
    pdf = BankingPDF()
    pdf.add_title_page()
    pdf.add_toc_page(toc_items)
    render_content(pdf, blocks, collect_toc=False)

    pdf.output(str(PDF_PATH))
    print(f"PDF saved: {PDF_PATH}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating DOCX...")
    build_docx()
    print()
    print("Generating PDF...")
    build_pdf()
    print()
    print("Done!")
